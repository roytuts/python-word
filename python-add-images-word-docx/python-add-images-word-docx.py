from docx import Document
from docx.shared import Inches

document = Document()

document.add_picture('computer.jpg', width=Inches(6))

document.add_page_break()

document.add_picture('sample.jpg')

#write to docx
document.save('images.docx')